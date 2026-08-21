"""Generate a concise v3 setup guide as a Word document.

Requires: pip install python-docx
"""

from pathlib import Path
import sys

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
except ImportError as exc:
    raise SystemExit("Install the guide dependency first: pip install python-docx") from exc


REPO = "https://github.com/martusha89/Memory"


def add_code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def add_steps(document: Document, steps: list[str]) -> None:
    for step in steps:
        document.add_paragraph(step, style="List Number")


def build_guide() -> Document:
    doc = Document()
    title = doc.add_heading("Memory v3 Setup Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph(
        "Persistent MCP memory on Cloudflare with versioned D1 records and a repairable Vectorize index."
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Before you begin", level=1)
    doc.add_paragraph(
        "You need Node.js 18 or later, a Cloudflare account, Wrangler access, and a strong secret. "
        "Cloudflare free tiers may be sufficient for personal use, but limits and pricing can change."
    )
    doc.add_paragraph(
        "D1 is canonical. Vectorize is a searchable projection. v3 keeps immutable memory versions and "
        "an indexing outbox so interrupted requests can be repaired. Automatic LLM consolidation is disabled."
    )

    doc.add_heading("1. Clone and install", level=1)
    add_code(doc, f"git clone {REPO}.git")
    add_code(doc, "cd Memory")
    add_code(doc, "npm install")

    doc.add_heading("2. Sign in and create resources", level=1)
    add_code(doc, "npx wrangler login")
    add_code(doc, "npx wrangler d1 create memory-db")
    add_code(doc, "npx wrangler vectorize create memory-index --dimensions=768 --metric=cosine")
    add_steps(
        doc,
        [
            "Copy the D1 database ID into wrangler.toml.",
            "If you changed resource names, update the D1 and Vectorize bindings in wrangler.toml.",
        ],
    )

    doc.add_heading("3. Initialize a fresh database", level=1)
    add_code(doc, "npx wrangler d1 execute memory-db --remote --file=schema.sql")
    doc.add_paragraph(
        "Do not create a partial schema by hand. schema.sql is the canonical fresh-install schema."
    )

    doc.add_heading("4. Upgrade an existing v2 database", level=1)
    doc.add_paragraph("Back up D1 first. Run only the migrations your database has not already applied.")
    add_code(
        doc,
        "npx wrangler d1 execute memory-db --remote --file=migrations/2026-06-12-pinned-and-archive.sql",
    )
    add_code(
        doc,
        "npx wrangler d1 execute memory-db --remote --file=migrations/2026-08-21-v3-safety-foundation.sql",
    )

    doc.add_heading("5. Set the required secret", level=1)
    add_code(doc, "npx wrangler secret put MEMORY_SECRET")
    doc.add_paragraph(
        "The server fails closed without MEMORY_SECRET. Use a long unique value and keep it in a password manager. "
        "Prefer Authorization headers. Secrets embedded in URLs can leak through history and logs."
    )

    doc.add_heading("6. Validate and deploy", level=1)
    add_code(doc, "npm run check")
    add_code(doc, "npm run deploy")
    doc.add_paragraph("Open https://YOUR-WORKER.workers.dev/health for version and service metadata.")
    doc.add_paragraph("Open the Worker root for the same-origin memory dashboard.")

    doc.add_heading("7. Connect an MCP client", level=1)
    doc.add_paragraph("For clients that support request headers, configure:")
    add_code(doc, "URL: https://YOUR-WORKER.workers.dev/mcp")
    add_code(doc, "Authorization: Bearer YOUR_SECRET")
    doc.add_paragraph(
        "Some MCP clients only accept a URL. The /mcp and /sse endpoints retain ?secret= compatibility, "
        "but treat those URLs as credentials and rotate the secret if they are exposed."
    )

    doc.add_heading("8. Enable index reconciliation", level=1)
    doc.add_paragraph("Uncomment the [triggers] section in wrangler.toml to retry projection failures every 15 minutes.")
    doc.add_paragraph(
        "You can also POST an authenticated request to /api/index/reconcile with {\"limit\": 50}."
    )

    doc.add_heading("Recovery notes", level=1)
    doc.add_paragraph(
        "A failed embedding or Vectorize request does not roll back the D1 memory. The outbox records the failure, "
        "lexical search can still find the D1 record, and reconciliation retries later. memory_versions preserves "
        "the previous content for inspection."
    )
    doc.add_paragraph(
        "The old automatic consolidation routine is intentionally disabled because a free-form model merge is not "
        "a safe replacement for canonical source history."
    )
    return doc


def main() -> None:
    output = Path(sys.argv[1]).expanduser() if len(sys.argv) > 1 else Path.home() / "Desktop" / "Memory v3 Setup Guide.docx"
    output.parent.mkdir(parents=True, exist_ok=True)
    build_guide().save(output)
    print(f"Saved to: {output}")


if __name__ == "__main__":
    main()
